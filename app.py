import openai
import nltk
from nltk.tokenize import word_tokenize
from fpdf import FPDF
from docx import Document
import streamlit as st
from langchain.text_splitter import RecursiveCharacterTextSplitter
import os 
# Initialize NLTK
nltk.download('punkt')

# Set OpenAI API key
openai.api_key = st.secrets["apikey"]

# Function to chunk large job descriptions using LangChain
def chunk_text(text, chunk_size=500, chunk_overlap=100):
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap
    )
    chunks = text_splitter.split_text(text)
    return chunks

# Function to extract keywords from each chunk using NLP
def extract_keywords_from_chunks(chunks):
    keywords = []
    for chunk in chunks:
        tokens = word_tokenize(chunk)
        chunk_keywords = [word for word in tokens if word.lower() in ['python', 'machine learning', 'deep learning', 'nlp', 'data wrangling']]
        keywords.extend(chunk_keywords)
    return keywords

# Function to generate professional summary using GPT-4 based on retrieved chunks
def generate_summary_from_chunks(chunks):
    summary = ""
    for chunk in chunks:
        prompt = f"Write a professional summary for a resume using the following information: {chunk}"
        response = openai.ChatCompletion.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "You are an expert in writing professional resumes also provide good quality resume and propery maintain line of sequence ."},
                {"role": "user", "content": prompt},
            ],
            max_tokens=200,
            temperature=0.0,
        )
        summary += response['choices'][0]['message']['content'].strip() + " "
    return summary.strip()

# Function to generate a PDF resume
def generate_pdf(user_data, summary, education, experience, skills, interests, projects):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", 'B', size=16)
    pdf.cell(200, 10, txt="Resume - " + user_data['name'], ln=True, align='C')
    
    pdf.set_font("Arial", size=12)
    pdf.cell(200, 10, txt="Contact Information", ln=True)
    pdf.multi_cell(0, 10, txt=f"Email: {user_data['email']}\nPhone: {user_data['phone']}\nLinkedIn: {user_data['linkedin']}\nGithub: {user_data['Github']}")

    # Add Professional Summary
    pdf.set_font("Arial", 'B', size=14)
    pdf.cell(200, 10, txt="Professional Summary", ln=True)
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(0, 10, txt=summary)

    # Add Education
    pdf.set_font("Arial", 'B', size=14)
    pdf.cell(200, 10, txt="Education", ln=True)
    pdf.set_font("Arial", size=12)
    for edu in education:
        pdf.multi_cell(0, 10, txt=f"{edu['degree']} in {edu['field']} - {edu['institution']}\nCGPA/Percentage: {edu['cgpa']} | Year of Passing: {edu['year']}\n")
    
    # Add Skills
    pdf.set_font("Arial", 'B', size=14)
    pdf.cell(200, 10, txt="Skills", ln=True)
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(0, 10, txt=", ".join(skills))
    
    # Add Projects
    if projects:
        pdf.set_font("Arial", 'B', size=14)
        pdf.cell(200, 10, txt="Projects", ln=True)
        pdf.set_font("Arial", size=12)
        for proj in projects:
            pdf.multi_cell(0, 10, txt=f"{proj['title']} | {proj['tech_stack']}\n{proj['description']}\n")
    
    # Add Experience
    if experience:
        pdf.set_font("Arial", 'B', size=14)
        pdf.cell(200, 10, txt="Experience", ln=True)
        pdf.set_font("Arial", size=12)
        for exp in experience:
            pdf.multi_cell(0, 10, txt=f"{exp['role']} at {exp['company']}\nDuration: {exp['duration']}\n")
            for point in exp['responsibilities']:
                pdf.multi_cell(0, 10, txt=f"- {point}")

    # Add Interests
    if interests:
        pdf.set_font("Arial", 'B', size=14)
        pdf.cell(200, 10, txt="Interests", ln=True)
        pdf.set_font("Arial", size=12)
        pdf.multi_cell(0, 10, txt=", ".join(interests))
    
    pdf_output = "generated_resume.pdf"
    pdf.output(pdf_output)
    return pdf_output

# Generate DOCX resume
def generate_docx(user_data, summary, education, experience, skills, interests, projects):
    doc = Document()
    doc.add_heading(f"Resume - {user_data['name']}", 0)
    doc.add_heading('Contact Information', level=1)
    doc.add_paragraph(f"Email: {user_data['email']}\nPhone: {user_data['phone']}\nLinkedIn: {user_data['linkedin']}\nGithub: {user_data['Github']}")

    doc.add_heading('Professional Summary', level=1)
    doc.add_paragraph(summary)

    doc.add_heading('Education', level=1)
    for edu in education:
        doc.add_paragraph(f"{edu['degree']} in {edu['field']} - {edu['institution']}")
        doc.add_paragraph(f"CGPA/Percentage: {edu['cgpa']} | Year of Passing: {edu['year']}")

    doc.add_heading('Skills', level=1)
    doc.add_paragraph(", ".join(skills))

    if experience:
        doc.add_heading('Experience', level=1)
        for exp in experience:
            doc.add_paragraph(f"{exp['role']} at {exp['company']}")
            doc.add_paragraph(f"Duration: {exp['duration']}")
            for point in exp['responsibilities']:
                doc.add_paragraph(f"- {point}")

    # Add Projects
    if projects:
        doc.add_heading('Projects', level=1)
        for proj in projects:
            doc.add_paragraph(f"{proj['title']} | {proj['tech_stack']}")
            doc.add_paragraph(f"{proj['description']}")

    doc.add_heading('Interests', level=1)
    doc.add_paragraph(", ".join(interests))

    doc_output = "generated_resume.docx"
    doc.save(doc_output)
    return doc_output

# Generate TXT resume
# def generate_txt(user_data, summary, education, experience, skills, interests, projects):
#     txt_output = "generated_resume.txt"
#     with open(txt_output, 'w') as file:
#         file.write(f"Resume - {user_data['name']}\n\n")
#         file.write("Contact Information:\n")
#         file.write(f"Email: {user_data['email']}\nPhone: {user_data['phone']}\nLinkedIn: {user_data['linkedin']}\n\n")
#         file.write("Professional Summary:\n")
#         file.write(summary + "\n\n")

#         file.write("Education:\n")
#         for edu in education:
#             file.write(f"{edu['degree']} in {edu['field']} - {edu['institution']}\n")
#             file.write(f"CGPA/Percentage: {edu['cgpa']} | Year of Passing: {edu['year']}\n\n")

#         file.write("Skills:\n")
#         file.write(", ".join(skills) + "\n\n")

#         if experience:
#             file.write("Experience:\n")
#             for exp in experience:
#                 file.write(f"{exp['role']} at {exp['company']}\n")
#                 file.write(f"Duration: {exp['duration']}\n")
#                 for point in exp['responsibilities']:
#                     file.write(f"- {point}\n")

#         if projects:
#             file.write("Projects:\n")
#             for proj in projects:
#                 file.write(f"{proj['title']} | {proj['tech_stack']}\n")
#                 file.write(f"{proj['description']}\n")

#         file.write("Interests:\n")
#         file.write(", ".join(interests) + "\n")

#     return txt_output

# Streamlit App to Collect User Data and Generate Resume
st.title("ATS-Friendly Resume Generator with RAG")

# Input Form for Resume Details
user_name = st.text_input("Full Name")
user_email = st.text_input("Email")
user_phone = st.text_input("Phone")
user_linkedin = st.text_input("LinkedIn Profile URL")
user_Github = st.text_input("Github Profile URL")
job_description = st.text_area("Job Description (from job post)")

# Education Section (Dynamic)
st.subheader("Education")
education_entries = []
num_education = st.number_input("Number of Education Entries", min_value=1, max_value=5, value=1)
for i in range(num_education):
    with st.expander(f"Education Entry {i+1}"):
        degree = st.text_input(f"Degree {i+1}")
        field = st.text_input(f"Field of Study {i+1}")
        institution = st.text_input(f"Institution {i+1}")
        cgpa = st.text_input(f"CGPA/Percentage {i+1}")
        year = st.text_input(f"Year of Completion {i+1}")
        education_entries.append({
            'degree': degree,
            'field': field,
            'institution': institution,
            'cgpa': cgpa,
            'year': year
        })

# Experience Section (Optional)
add_experience = st.checkbox("Add Work Experience")
experience_entries = []
if add_experience:
    num_experience = st.number_input("Number of Work Experiences", min_value=1, max_value=5, value=1)
    for i in range(num_experience):
        with st.expander(f"Experience Entry {i+1}"):
            role = st.text_input(f"Role {i+1}")
            company = st.text_input(f"Company {i+1}")
            duration = st.text_input(f"Duration {i+1}")
            responsibilities = st.text_area(f"Responsibilities {i+1} (comma separated)").split(",")
            experience_entries.append({
                'role': role,
                'company': company,
                'duration': duration,
                'responsibilities': responsibilities
            })

# Skills Section (Required)
st.subheader("Skills")
skills = st.text_area("Skills (comma separated)").split(",")

# Projects Section (Optional)
add_projects = st.checkbox("Add Projects")
projects = []
if add_projects:
    num_projects = st.number_input("Number of Projects", min_value=1, max_value=5, value=1)
    for i in range(num_projects):
        with st.expander(f"Project {i+1}"):
            title = st.text_input(f"Project Title {i+1}")
            tech_stack = st.text_input(f"Tech Stack {i+1}")
            description = st.text_area(f"Description {i+1}")
            projects.append({
                'title': title,
                'tech_stack': tech_stack,
                'description': description
            })

# Interests Section (Optional)
add_interests = st.checkbox("Add Interests")
interests = []
if add_interests:
    interests = st.text_area("Interests (comma separated)").split(",")

# Process Job Description and Generate Summary
if st.button("Generate Resume"):
    if user_name and user_email and user_phone and user_linkedin and user_Github and job_description:
        # Chunk job description
        chunks = chunk_text(job_description)

        # Extract keywords from chunks
        keywords = extract_keywords_from_chunks(chunks)

        # Generate professional summary from job description chunks
        professional_summary = generate_summary_from_chunks(chunks)

        # Prepare user data for PDF, DOCX, and TXT
        user_data = {
            'name': user_name,
            'email': user_email,
            'phone': user_phone,
            'linkedin': user_linkedin,
            'Github' : user_Github
        }

        # Generate Resume Outputs
        pdf_resume = generate_pdf(user_data, professional_summary, education_entries, experience_entries, skills, interests, projects)
        docx_resume = generate_docx(user_data, professional_summary, education_entries, experience_entries, skills, interests, projects)
        #txt_resume = generate_txt(user_data, professional_summary, education_entries, experience_entries, skills, interests, projects)

        # Display Download Links for Resumes
        st.success("Resume Generated Successfully!")
        st.download_button("Download PDF Resume", data=open(pdf_resume, "rb"), file_name=pdf_resume)
        st.download_button("Download DOCX Resume", data=open(docx_resume, "rb"), file_name=docx_resume)
        #st.download_button("Download TXT Resume", data=open(txt_resume, "rb"), file_name=txt_resume)

    else:
        st.error("Please fill out all required fields!")
